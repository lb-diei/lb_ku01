
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
公文格式化工具 - 修复版本
按照公文管理规范自动排版Word文档，包含首行缩进和页码功能
"""

import os
import sys
import argparse
import glob
from pathlib import Path
from datetime import datetime
import re

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
    from docx.oxml.shared import OxmlElement
except ImportError as e:
    print("错误：缺少必要的依赖包，请先安装 python-docx")
    print("安装命令：pip install python-docx")
    sys.exit(1)

class OfficialDocumentFormatter:
    """公文格式化类"""
    
    def __init__(self):
        self.format_settings = {
            'page': {
                'height': Cm(29.7),        # A4高度29.7cm
                'width': Cm(21),           # A4宽度21cm
                'top_margin': Cm(3.7),     # 上边距3.7cm
                'bottom_margin': Cm(3.2),  # 下边距3.2cm
                'left_margin': Cm(2.8),    # 左边距2.8cm
                'right_margin': Cm(2.6),   # 右边距2.6cm
            },
            'fonts': {
                'title': {'size': Pt(22), 'name': '方正小标宋简体'},                      # 二号
                'level1': {'size': Pt(16), 'name': '黑体', 'bold': True},                 # 三号黑体
                'level2': {'size': Pt(16), 'name': '楷体_GB2312', 'bold': True},          # 三号楷体加粗
                'level3': {'size': Pt(16), 'name': '仿宋_GB2312', 'bold': True},          # 三号仿宋加粗
                'body': {'size': Pt(16), 'name': '仿宋_GB2312', 'bold': False},           # 三号仿宋
                'footer': {'size': Pt(14), 'name': 'Times New Roman', 'bold': False},     # 四号
            },
            'spacing': {
                'line_spacing': Pt(28),        # 行间距28磅
                'title_spacing': Pt(0),       # 标题段间距30磅
                'body_spacing': Pt(0),         # 正文段间距
                'first_line_indent': Cm(1.7),  # 首行缩进2个字符（约0.85cm）
            }
        }
    
    def detect_title_level(self, text):
        """检测标题级别"""
        text = text.strip()
        if text.startswith(('一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、')):
            return 'level1'
        elif text.startswith(('（一）', '（二）', '（三）', '（四）', '（五）', '（六）', '（七）', '（八）', '（九）', '（十）')):
            return 'level2'
        elif text.startswith(tuple(f"{i}." for i in range(1, 20))):
            return 'level3'
        elif text.startswith(('（1）', '（2）', '（3）', '（4）', '（5）', '（6）', '（7）', '（8）', '（9）', '（10）')):
            return 'level3'
        else:
            return 'body'
    
    def apply_font_formatting(self, run, font_type):
        """应用字体格式"""
        settings = self.format_settings['fonts'][font_type]
        run.font.size = settings['size']
        run.font.bold = settings.get('bold', False)
        
        # 设置字体
        run.font.name = settings['name']
        if '方正' in settings['name'] or '黑体' in settings['name'] or '楷体_GB2312' in settings['name'] or '仿宋_GB2312' in settings['name']:
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), settings['name'])
    
    def setup_page_layout(self, section):
        """设置页面布局"""
        page_settings = self.format_settings['page']
        section.page_height = page_settings['height']
        section.page_width = page_settings['width']
        section.top_margin = page_settings['top_margin']
        section.bottom_margin = page_settings['bottom_margin']
        section.left_margin = page_settings['left_margin']
        section.right_margin = page_settings['right_margin']
    
    def create_page_number_element(self, paragraph):
        """创建页码元素"""
        # 创建页码字段
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        instrText.set(qn('w:space'), 'preserve')
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        # 添加分隔符
        run = paragraph.add_run()
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        
        return run
    
    def add_page_numbers(self, doc):
        """添加页码到文档的页脚"""
        try:
            for i, section in enumerate(doc.sections):
                footer = section.footer
                
                # 清除现有的页脚内容
                if footer.paragraphs:
                    for para in footer.paragraphs:
                        p = para._element
                        p.getparent().remove(p)
                
                # 创建页脚段落
                footer_para = footer.add_paragraph()
                footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 添加左括号
                left_run = footer_para.add_run("-")
                left_run.font.size = self.format_settings['fonts']['footer']['size']
                left_run.font.name = self.format_settings['fonts']['footer']['name']
                
                # 添加页码字段
                page_run = self.create_page_number_element(footer_para)
                page_run.font.size = self.format_settings['fonts']['footer']['size']
                page_run.font.name = self.format_settings['fonts']['footer']['name']
                
                # 添加右括号
                right_run = footer_para.add_run("-")
                right_run.font.size = self.format_settings['fonts']['footer']['size']
                right_run.font.name = self.format_settings['fonts']['footer']['name']
                
                print(f"已为第 {i+1} 节添加页码")
                
        except Exception as e:
            print(f"添加页码时出错: {str(e)}")
            # 备选方案：手动添加页码
            self.add_manual_page_numbers(doc)
    
    def add_manual_page_numbers(self, doc):
        """手动添加页码（备选方案）"""
        for i, section in enumerate(doc.sections):
            footer = section.footer
            if footer.paragraphs:
                for para in footer.paragraphs:
                    p = para._element
                    p.getparent().remove(p)
            
            footer_para = footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
            # 添加页码格式 (-1-)
            page_number = i + 1
            footer_run = footer_para.add_run(f"- {page_number} -")
            footer_run.font.size = self.format_settings['fonts']['footer']['size']
            footer_run.font.name = self.format_settings['fonts']['footer']['name']
            print(f"使用备选方案添加页码: - {page_number} -")
    
    def process_document_structure(self, doc, new_doc):
        """处理文档结构"""
        title_processed = False
        title_text = ""
        content_paragraphs = []


        
        # 收集所有有内容的段落
        for para in doc.paragraphs:
            if para.text.strip():
                content_paragraphs.append(para)

         # 遍历所有段落，设置段后间距为0
        for paragraph in doc.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(0)
        
        if content_paragraphs:
            # 第一个段落作为标题
            title_text = content_paragraphs[0].text.strip()
            title_para = new_doc.add_paragraph()
            title_run = title_para.add_run(title_text)
            
            # 设置标题格式
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.apply_font_formatting(title_run, 'title')
            title_para.paragraph_format.space_after = self.format_settings['spacing']['title_spacing']
            title_processed = True
        
        # 添加空行（正文前空一行）
        # new_doc.add_paragraph()
        # new_para1 = new_doc.add_paragraph()
        # new_para1.paragraph_format.space_after = self.format_settings['spacing']['line_spacing']
        
        # 处理正文内容（跳过标题段落）
        for para in content_paragraphs[1:] if title_processed else content_paragraphs:
            if not para.text.strip():
                continue
                
            new_para = new_doc.add_paragraph()
            text_content = para.text.strip()
            
            # 检测标题级别并应用格式
            level = self.detect_title_level(text_content)
            run = new_para.add_run(text_content)
            self.apply_font_formatting(run, level)
            
            # 设置段落格式
            new_para.paragraph_format.line_spacing = self.format_settings['spacing']['line_spacing']
            
            # 为正文段落添加首行缩进2字符
            if level == 'body':
                new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                new_para.paragraph_format.first_line_indent = Pt(32)
            else:
                new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # 标题段落也缩进
                new_para.paragraph_format.first_line_indent = Pt(32)
    
    def add_signature_block(self, doc, organization_name):
        """添加落款 - 修复版本"""
        # 与正文空两行
        # doc.add_paragraph()
        # doc.add_paragraph()
        
        # 添加单位名称
        org_para = doc.add_paragraph()
        org_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        org_run = org_para.add_run(organization_name)
        self.apply_font_formatting(org_run, 'body')
        
    
    def format_document(self, input_path, output_path, organization_name="发文单位名称年月日"):
        """格式化文档主函数"""
        try:
            # 检查输入文件是否存在
            if not os.path.exists(input_path):
                print(f"错误：输入文件不存在 - {input_path}")
            
            print(f"正在读取文档: {input_path}")
            doc = Document(input_path)
            
            # 创建新文档
            new_doc = Document()
            
            # 设置页面布局
            section = new_doc.sections[0]
            self.setup_page_layout(section)
            
            # 处理文档结构
            self.process_document_structure(doc, new_doc)
            
            # 添加落款
            self.add_signature_block(new_doc, organization_name)
            
            # 添加页码
            self.add_page_numbers(new_doc)
            
            # 保存文档
            print(f"正在保存格式化后的文档: {output_path}")
            new_doc.save(output_path)
            
            # 验证页码是否添加成功
            self.verify_page_numbers(output_path)
            
            print(f"文档格式化完成！输出文件: {output_path}")
            return True
            
        except Exception as e:
            print(f"格式化过程中出现错误: {str(e)}")
            import traceback
            traceback.print_exc()
            return False
    
    def verify_page_numbers(self, file_path):
        """验证页码是否添加成功 - 修复版本"""
        try:
            test_doc = Document(file_path)
            has_page_numbers = False
            
            for section in test_doc.sections:
                footer = section.footer
                if footer and footer.paragraphs:
                    for para in footer.paragraphs:
                        if para.text and '-' in para.text:
                            has_page_numbers = True
                            break
                
                if has_page_numbers:
                    break
            
            if has_page_numbers:
                print("✓ 页码添加成功")
            else:
                print("⚠ 页码可能未正确添加，请手动检查")
                
        except Exception as e:
            print(f"验证页码时出错: {str(e)}")
            return False

def get_word_files():
    """获取当前目录下的Word文档"""
    word_files = []
    for pattern in ["*.docx", "*.doc"]:
        word_files.extend(glob.glob(pattern))
    return [f for f in word_files if not f.endswith("_格式化.docx") and not f.endswith("_格式化.doc")]

def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='公文格式化工具 - 修复版本')
    parser.add_argument('input', nargs='?', help='输入Word文档路径')
    parser.add_argument('output', nargs='?', help='输出Word文档路径')
    parser.add_argument('--organization', '-o', default='发文单位名称年月日', help='发文单位名称')
    
    args = parser.parse_args()
    formatter = OfficialDocumentFormatter()
    
    # 处理输入输出路径
    if args.input:
        input_path = args.input
    else:
        # 如果没有指定输入文件，查找当前目录下的Word文档
        word_files = get_word_files()
        if not word_files:
            print("错误：当前目录下未找到Word文档（.docx 或 .doc）")
            print("请指定输入文件路径，或将Word文档放在当前目录")
            sys.exit(1)
        
        # 让用户选择文件
        print("当前目录下找到以下Word文档：")
        for i, file in enumerate(word_files, 1):
            print(f"{i}. {file}")
        
        try:
            choice = int(input("请选择要格式化的文档编号: ")) - 1
            if 0 <= choice < len(word_files):
                input_path = word_files[choice]
            else:
                print("错误：选择无效")
                sys.exit(1)
        except ValueError:
            print("错误：请输入有效的数字")
            sys.exit(1)
    
    if args.output:
        output_path = args.output
    else:
        # 生成默认输出文件名
        input_file = Path(input_path)
        output_path = input_file.stem + "_格式化.docx"
    
    # 执行格式化
    success = formatter.format_document(input_path, output_path, args.organization)
    if not success:
        sys.exit(1)

if __name__ == "__main__":
    main()

#<code_start project_name=公文格式化工具 filename=requirements.txt title=项目依赖文件 entrypoint=false runnable=false project_final_file=true>
#python-docx>=0.8.11
#<code_end>

